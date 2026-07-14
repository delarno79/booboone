"""Convert the client's keyword .docx into keywords.json.

Usage:
    python convert_docx.py "Booboone Keyword List.docx"
    python convert_docx.py "Booboone Keyword List.docx" -o keywords.json

The parser is heuristic because keyword documents vary. It detects categories
from headings / bold / ALL-CAPS lines and from table headers, and treats the
remaining lines / bullets as keywords. It ALWAYS prints a summary so you can
sanity-check the result before publishing. If the auto-detection is off, you
can hand-edit keywords.json afterwards — the format is simple:

    {
      "Business & Finance": ["small business bank account", ...],
      "Beauty & Fashion":   ["eye creams for dark circles", ...]
    }
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)


UNCATEGORIZED = "Uncategorized"


def _looks_like_category(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 60:
        return False
    style = (paragraph.style.name or "").lower() if paragraph.style else ""
    if style.startswith("heading") or style == "title":
        return True
    # All runs bold?
    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs and all(r.bold for r in runs):
        return True
    # ALL CAPS with few words (e.g. "BUSINESS & FINANCE").
    letters = re.sub(r"[^A-Za-z]", "", text)
    if letters and text.upper() == text and len(text.split()) <= 5:
        return True
    return False


def _clean_keyword(text: str) -> str:
    text = text.strip()
    # Strip common bullet / numbering prefixes.
    text = re.sub(r"^\s*(?:[-*•·▪◦]|\d+[.)])\s*", "", text)
    return text.strip(" \t.-•·")


def parse_docx(path: Path) -> dict[str, list[str]]:
    if path.suffix.lower() == ".doc":
        raise ValueError(
            "This is an old .doc file. Please open it in Word and 'Save As' .docx first."
        )
    document = Document(str(path))
    categories: dict[str, list[str]] = {}
    current = UNCATEGORIZED

    # 1) Paragraphs (headings + bullet lists).
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _looks_like_category(para):
            current = text
            categories.setdefault(current, [])
        else:
            kw = _clean_keyword(text)
            if kw:
                categories.setdefault(current, []).append(kw)

    # 2) Tables (header row = category, column cells = keywords).
    for table in document.tables:
        rows = table.rows
        if not rows:
            continue
        headers = [c.text.strip() for c in rows[0].cells]
        for col_idx, header in enumerate(headers):
            cat = header or UNCATEGORIZED
            categories.setdefault(cat, [])
            for row in rows[1:]:
                if col_idx < len(row.cells):
                    kw = _clean_keyword(row.cells[col_idx].text)
                    if kw:
                        categories[cat].append(kw)

    # Drop empty categories and de-duplicate.
    cleaned: dict[str, list[str]] = {}
    for cat, words in categories.items():
        seen: set[str] = set()
        deduped = []
        for w in words:
            key = w.lower()
            if w and key not in seen:
                seen.add(key)
                deduped.append(w)
        if deduped:
            cleaned[cat] = deduped
    return cleaned


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert keyword .docx to JSON")
    parser.add_argument("docx", help="Path to the .docx keyword file")
    parser.add_argument("-o", "--output", default="keywords.json")
    args = parser.parse_args()

    path = Path(args.docx)
    if not path.exists():
        print(f"File not found: {path}")
        return 1

    try:
        categories = parse_docx(path)
    except ValueError as err:
        print(err)
        return 1

    if not categories:
        print("No keywords found. Check the document, or build keywords.json by hand.")
        return 1

    out = Path(args.output)
    out.write_text(json.dumps(categories, indent=2, ensure_ascii=False), encoding="utf-8")

    total = sum(len(v) for v in categories.values())
    print(f"\nWrote {out} — {len(categories)} categories, {total} keywords:\n")
    for cat, words in categories.items():
        print(f"  {cat}: {len(words)} keywords")
        for w in words[:3]:
            print(f"      - {w}")
        if len(words) > 3:
            print(f"      ... (+{len(words) - 3} more)")
    print("\nReview keywords.json and edit by hand if any category looks wrong.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
